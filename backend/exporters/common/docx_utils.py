from __future__ import annotations

import io
import base64
import struct
from typing import Optional

import pandas as pd

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    HAS_DOCX = True
except ImportError:  # pragma: no cover
    HAS_DOCX = False
    Document = None  # type: ignore[assignment]
    WD_ORIENT = None  # type: ignore[assignment]
    WD_SECTION = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]
    Inches = None  # type: ignore[assignment]
    RGBColor = None  # type: ignore[assignment]


THEME_RGB = RGBColor(31, 78, 121) if HAS_DOCX else None
EMU_PER_INCH = 914400
WIDE_FIGURE_MIN_WIDTH_PX = 1400
WIDE_FIGURE_MIN_ASPECT_RATIO = 1.65


def create_document(
    font_name: str = "Microsoft YaHei",
    east_asia_font: Optional[str] = None,
    font_size: float = 10.5,
):
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not available")

    east_asia = east_asia_font or font_name
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = font_name
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    styles["Normal"].font.size = Pt(font_size)
    return doc


def add_report_title(doc, title: str, subtitle: Optional[str] = None) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)
    if THEME_RGB is not None:
        title_run.font.color.rgb = THEME_RGB

    if subtitle:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(12)

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider.add_run("────────────────────────")
    divider_run.font.size = Pt(10)
    if THEME_RGB is not None:
        divider_run.font.color.rgb = THEME_RGB


def add_heading(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    if THEME_RGB is not None:
        run.font.color.rgb = THEME_RGB


def add_report_note(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    if THEME_RGB is not None:
        run.font.color.rgb = THEME_RGB


def style_table_header_row(table) -> None:
    if not table.rows:
        return

    for cell in table.rows[0].cells:
        if not cell.paragraphs:
            continue
        for run in cell.paragraphs[0].runs:
            run.bold = True
            if THEME_RGB is not None:
                run.font.color.rgb = THEME_RGB


def add_df_table(doc, frame: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(frame.columns):
        header_paragraph = table.rows[0].cells[idx].paragraphs[0]
        header_paragraph.add_run(str(column))
    style_table_header_row(table)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def _png_dimensions(png_bytes: bytes) -> Optional[tuple[int, int]]:
    if len(png_bytes) < 24 or not png_bytes.startswith(b"\x89PNG\r\n\x1a\n") or png_bytes[12:16] != b"IHDR":
        return None
    width, height = struct.unpack(">II", png_bytes[16:24])
    return width, height


def _content_width_inches(section) -> float:
    return max(1.0, float(section.page_width - section.left_margin - section.right_margin) / EMU_PER_INCH)


def _copy_section_margins(target, source) -> None:
    target.left_margin = source.left_margin
    target.right_margin = source.right_margin
    target.top_margin = source.top_margin
    target.bottom_margin = source.bottom_margin
    target.header_distance = source.header_distance
    target.footer_distance = source.footer_distance


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    width, height = section.page_width, section.page_height
    section.page_width = max(width, height)
    section.page_height = min(width, height)


def _set_portrait(section, source) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    width, height = source.page_width, source.page_height
    section.page_width = min(width, height)
    section.page_height = max(width, height)
    _copy_section_margins(section, source)


def _should_use_landscape_figure(png_bytes: bytes) -> bool:
    dimensions = _png_dimensions(png_bytes)
    if not dimensions:
        return False
    width, height = dimensions
    return width >= WIDE_FIGURE_MIN_WIDTH_PX and width / max(height, 1) >= WIDE_FIGURE_MIN_ASPECT_RATIO


def _add_picture_paragraph(doc, png_bytes: bytes, width_inches: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(width_inches))


def add_png_figure(doc, png_bytes: bytes, caption: str, width_inches: float = 6.2) -> None:
    if not png_bytes:
        return
    if _should_use_landscape_figure(png_bytes):
        source_section = doc.sections[-1]
        landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
        _copy_section_margins(landscape_section, source_section)
        _set_landscape(landscape_section)
        _add_picture_paragraph(doc, png_bytes, _content_width_inches(landscape_section))
    else:
        source_section = None
        _add_picture_paragraph(doc, png_bytes, width_inches)

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    if THEME_RGB is not None:
        caption_run.font.color.rgb = THEME_RGB
    if source_section is not None:
        portrait_section = doc.add_section(WD_SECTION.CONTINUOUS)
        _set_portrait(portrait_section, source_section)


def png_from_report_images(report_images: Optional[dict], key: str) -> bytes:
    if not report_images:
        return b""
    value = report_images.get(key)
    if not isinstance(value, str) or not value:
        return b""
    payload = value.split(",", 1)[1] if value.startswith("data:") and "," in value else value
    try:
        return base64.b64decode(payload, validate=True)
    except Exception:
        return b""
