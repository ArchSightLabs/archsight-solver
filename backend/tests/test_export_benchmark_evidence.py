from __future__ import annotations

import io

import pandas as pd
import pytest
from docx import Document

from backend.benchmarks.catalog import find_benchmark_case
from backend.examples.public_validation_projects import build_public_validation_projects
from backend.contracts.json_schemas import API_SCHEMA_VERSION
from backend.exporters.common.evidence import _beam_vertical_load_kn, _structure_load_resultant, build_evidence_tables, build_report_review_table
from backend.application.truss_analysis import _loads_for_combination as truss_loads_for_combination
from backend.services.export_service import build_report_model, export_report
from backend.tests.test_beam_workbench import beam_payload
from backend.tests.test_frame_workbench import frame_payload, frame_payload_with_combination_tags
from backend.tests.test_truss_workbench import _base_payload as truss_payload


def _table_text(table) -> str:
    return "\n".join(" | ".join(row) for row in table.astype(str).values.tolist())


def _docx_text(artifact) -> str:
    docx = Document(io.BytesIO(artifact.buffer.getvalue()))
    return "\n".join(
        [paragraph.text for paragraph in docx.paragraphs]
        + [cell.text for table in docx.tables for row in table.rows for cell in row.cells]
    )


def _xlsx_text(artifact) -> str:
    with pd.ExcelFile(io.BytesIO(artifact.buffer.getvalue())) as xls:
        return "\n".join(
            pd.read_excel(xls, sheet_name=sheet_name, header=None).astype(str).to_string()
            for sheet_name in xls.sheet_names
        )


def test_selected_frame_combination_projects_the_entire_report_body():
    payload = frame_payload_with_combination_tags()
    report = build_report_model(
        {
            **payload,
            "resultSource": {
                "source": "combination",
                "id": "ULS1",
                "label": "承载能力基本组合",
                "description": "1.2DL + 1.5WL",
            },
        },
        analysis_type="frame",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
    )
    selected = next(item for item in report["loadCombinationResults"] if item["id"] == "ULS1")

    assert report["summary"] == selected["summary"]
    assert report["nodeResults"] == selected["nodeResults"]
    assert report["memberResults"] == selected["memberResults"]
    assert report["memberDiagrams"] == selected["memberDiagrams"]
    selected_loads = report["structure"]["loads"]
    assert selected_loads[0]["type"] == "distributed"
    assert selected_loads[0]["qStartKnPerM"] == pytest.approx(-21.6)
    assert selected_loads[1]["type"] == "nodal"
    assert selected_loads[1]["fxKn"] == pytest.approx(36.0)
    equilibrium = build_evidence_tables(report, "frame", "测试材料")["校核证据"]
    equilibrium_text = _table_text(equilibrium)
    assert equilibrium_text.count("相对误差 0.0%") >= 2
    assert report["selectedResult"]["factors"] == {"DL": 1.2, "WL": 1.5}
    assert "ULS1" in _docx_text(export_report(report, "docx"))
    assert "ULS1" in _xlsx_text(export_report(report, "xlsx"))


def test_selected_beam_case_and_combination_project_resolved_input_and_equilibrium():
    payload = beam_payload()
    payload["loadCases"] = [
        {"id": "DL", "title": "恒载", "loads": [{"type": "uniform", "qKnPerM": 8.0}]},
        {"id": "LL", "title": "活载", "loads": [{"type": "point", "magnitudeKn": 20.0, "x": 2.5}]},
    ]
    payload["loadCombinations"] = [
        {"id": "ULS1", "title": "基本组合", "factors": {"DL": 1.2, "LL": 1.4}},
    ]

    case_report = build_report_model(
        {**payload, "resultSource": {"source": "case", "id": "LL", "label": "活载"}},
        analysis_type="beam",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
    )
    assert case_report["request"]["load_type"] == "point"
    assert case_report["request"]["point_load_kn"] == pytest.approx(20.0)
    assert case_report["status"] == case_report["summary"]["status"]
    assert case_report["max_deflection_position_m"] == pytest.approx(case_report["summary"]["maxDeflectionPositionM"])
    case_docx = _docx_text(export_report(case_report, "docx"))
    assert "集中荷载 P" in case_docx and "20.0 kN" in case_docx
    assert "均布荷载 q | 10.0 kN/m" not in case_docx

    combination_report = build_report_model(
        {**payload, "resultSource": {"source": "combination", "id": "ULS1", "label": "基本组合"}},
        analysis_type="beam",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
    )
    assert combination_report["request"]["load_type"] == "combination"
    assert combination_report["request"]["selected_load_combination"]["factors"] == {"DL": 1.2, "LL": 1.4}
    assert combination_report["status"] == combination_report["summary"]["status"]
    assert combination_report["max_deflection_position_m"] == pytest.approx(
        combination_report["summary"]["maxDeflectionPositionM"]
    )
    combination_docx = _docx_text(export_report(combination_report, "docx"))
    assert "1.2×DL + 1.4×LL" in combination_docx
    assert "均布荷载 q | 10.0 kN/m" not in combination_docx
    evidence_text = _table_text(build_evidence_tables(combination_report, "beam", "测试材料")["校核证据"])
    assert evidence_text.count("相对误差 0.0%") >= 1


def test_selected_truss_case_and_combination_project_resolved_loads():
    payload = truss_payload()
    payload["structure"]["loads"] = []
    payload["structure"]["loadCases"] = [
        {"id": "VL", "title": "竖向", "loads": [{"type": "nodal", "node": "N3", "fxKn": 0.0, "fyKn": -12.0}]},
        {"id": "HL", "title": "水平", "loads": [{"type": "nodal", "node": "N4", "fxKn": 8.0, "fyKn": 0.0}]},
    ]
    payload["structure"]["loadCombinations"] = [
        {"id": "COMB1", "title": "组合", "factors": {"VL": 1.2, "HL": 1.0}},
    ]

    case_report = build_report_model(
        {**payload, "resultSource": {"source": "case", "id": "VL", "label": "竖向"}},
        analysis_type="truss",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
    )
    assert case_report["structure"]["loads"] == [
        {"type": "nodal", "node": "N3", "fxKn": 0.0, "fyKn": -12.0}
    ]

    combination_report = build_report_model(
        {**payload, "resultSource": {"source": "combination", "id": "COMB1", "label": "组合"}},
        analysis_type="truss",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
    )
    combination_loads = combination_report["structure"]["loads"]
    assert combination_loads[0]["node"] == "N3"
    assert combination_loads[0]["fyKn"] == pytest.approx(-14.4)
    assert combination_loads[1]["node"] == "N4"
    assert combination_loads[1]["fxKn"] == pytest.approx(8.0)
    evidence_text = _table_text(build_evidence_tables(combination_report, "truss", "测试材料")["校核证据"])
    assert "残差 0.0 kN" in evidence_text
    assert "COMB1" in _docx_text(export_report(combination_report, "docx"))
    assert "COMB1" in _xlsx_text(export_report(combination_report, "xlsx"))


def test_report_equilibrium_helpers_cover_partial_and_member_load_contracts():
    assert _beam_vertical_load_kn(
        {
            "load_type": "uniform",
            "q_kn": 10.0,
            "total_length": 10.0,
            "uniform_start": 2.5,
            "uniform_end": 7.5,
        }
    ) == pytest.approx(50.0)

    structure = {
        "nodes": [
            {"id": "N1", "x": 0.0, "y": 0.0},
            {"id": "N2", "x": 4.0, "y": 0.0},
        ],
        "members": [{"id": "M1", "start": "N1", "end": "N2"}],
        "loads": [
            {"type": "member_point", "member": "M1", "direction": "global_y", "forceKn": -10.0},
            {
                "type": "distributed",
                "member": "M1",
                "direction": "local_y",
                "qStartKnPerM": -2.0,
                "qEndKnPerM": -2.0,
                "startRatio": 0.25,
                "endRatio": 0.75,
            },
        ],
    }
    resultant = _structure_load_resultant(structure)
    assert resultant["fx"] == pytest.approx(0.0)
    assert resultant["fy"] == pytest.approx(-14.0)

    temperature_loads = truss_loads_for_combination(
        {"id": "T-COMB", "factors": {"T": 1.5}},
        [
            {
                "id": "T",
                "loads": [
                    {"type": "temperature", "member": "M1", "deltaTempC": 20.0, "alphaPerC": 1.2e-5}
                ],
            }
        ],
    )
    assert temperature_loads == [
        {"type": "temperature", "member": "M1", "deltaTempC": 30.0, "alphaPerC": 1.2e-5}
    ]


def test_export_evidence_tables_include_public_benchmark_source_and_expected_values():
    case = find_benchmark_case("truss-simple-roof")
    examples = build_public_validation_projects()
    benchmark = next(
        obj["benchmark"]
        for project in examples["projects"]
        for obj in project["project"]["objects"]
        if obj["benchmark"]["caseId"] == "truss-simple-roof"
    )
    payload = {**case["payload"], "benchmark": benchmark}

    report = build_report_model(
        payload,
        analysis_type="truss",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
    )
    evidence = build_evidence_tables(report, "truss", "测试材料")["校核证据"]
    rows = evidence.astype(str).values.tolist()

    assert any(row[0] == "当前算例验证等级" and row[1] == "B 级验证" for row in rows)
    assert any(row[0] == "当前算例来源" and "truss-simple-roof" in row[1] for row in rows)
    assert any(row[0] == "当前算例标准值" and "标准值：" in row[1] for row in rows)
    assert any(row[0] == "当前算例容许误差" and "容许误差：" in row[1] for row in rows)
    assert all("弯矩" not in row[1] and "剪力" not in row[1] for row in rows if row[0].startswith("当前算例"))


def test_export_evidence_resolves_learning_review_from_catalog_instead_of_client_labels():
    case = find_benchmark_case("BM-009")
    examples = build_public_validation_projects()
    benchmark = next(
        obj["benchmark"]
        for project in examples["projects"]
        for obj in project["project"]["objects"]
        if obj["benchmark"]["caseId"] == "BM-009"
    )
    payload = {
        **case["payload"],
        "benchmark": benchmark,
        "learningReview": {
            "schemaVersion": 1,
            "pathId": "truss-force-path",
            "caseId": "BM-009",
            "reviewed": True,
            "answers": [
                {"predictionId": "support-reactions", "selectedOptionId": "symmetric-30"},
                {"predictionId": "member-forces", "selectedOptionId": "diagonal-compression-bottom-tension"},
                {"predictionId": "node-displacement", "selectedOptionId": "apex-down-roller-right"},
            ],
            "clientLabel": "不得进入计算书的伪造文案",
        },
    }

    report = build_report_model(
        payload,
        analysis_type="truss",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
    )
    rows = build_evidence_tables(report, "truss", "测试材料")["校核证据"].astype(str).values.tolist()
    text = "\n".join(" | ".join(row) for row in rows)

    assert "三杆桁架：先判拉压，再看位移" in text
    assert "支座反力" in text
    assert "判断一致" in text
    assert "不得进入计算书的伪造文案" not in text

    docx_artifact = export_report(report, "docx")
    docx = Document(docx_artifact.buffer)
    docx_text = "\n".join(
        cell.text
        for table in docx.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "三杆桁架：先判拉压，再看位移" in docx_text
    assert "判断一致" in docx_text

    xlsx_artifact = export_report(report, "xlsx")
    xlsx_text = pd.read_excel(xlsx_artifact.buffer, sheet_name="05_校核证据", header=None).to_string()
    assert "三杆桁架：先判拉压，再看位移" in xlsx_text
    assert "判断一致" in xlsx_text


def test_frame_truss_evidence_input_summary_preserves_material_and_support_semantics():
    frame_solution = {
        "structure": {
            "nodes": [
                {"id": "N1", "x": 0, "y": 0, "supportType": "fixed"},
                {"id": "N2", "x": 4, "y": 0, "supportType": "roller", "supportAngleDeg": 45},
            ],
            "members": [
                {"id": "C1", "start": "N1", "end": "N2", "materialId": "q235", "E_GPa": 206, "A_cm2": 120, "I_cm4": 8000},
                {"id": "B1", "start": "N1", "end": "N2", "materialId": "custom", "E_GPa": 198.5, "A_cm2": 100, "I_cm4": 6000},
            ],
            "loads": [],
        },
        "summary": {"allowableMm": 20},
        "nodeResults": [],
        "memberResults": [],
    }
    truss_solution = {
        "structure": {
            "nodes": [
                {"id": "N1", "x": 0, "y": 0, "supportType": "pinned"},
                {"id": "N2", "x": 4, "y": 0, "supportType": "roller"},
            ],
            "members": [
                {"id": "M1", "start": "N1", "end": "N2", "materialId": "q345", "E_GPa": 210, "A_cm2": 24},
            ],
            "loads": [],
        },
        "summary": {"allowableMm": 10},
        "nodeResults": [],
        "memberResults": [],
    }

    frame_text = _table_text(build_evidence_tables(frame_solution, "frame", "测试材料")["工程输入摘要"])
    truss_text = _table_text(build_evidence_tables(truss_solution, "truss", "测试材料")["工程输入摘要"])

    assert "材料适用范围 | 材料名称为项目默认材料说明；框架整体刚度按各构件 E_GPa / A_cm2 / I_cm4 输入装配。" in frame_text
    assert "构件弹性模量分布 | " in frame_text
    assert "Q235 · E=206 GPa：1 个构件" in frame_text
    assert "E=198.5 GPa：1 个构件" in frame_text
    assert "支座体系说明 | 平面框架节点自由度为 ux、uy、rz" in frame_text

    assert "材料适用范围 | 材料名称为项目默认材料说明；桁架整体刚度按各杆件 E_GPa / A_cm2 输入装配。" in truss_text
    assert "杆件弹性模量分布 | Q345 · E=210 GPa：1 个杆件" in truss_text
    assert "支座体系说明 | 平面桁架节点仅含 ux、uy 平动自由度" in truss_text


def test_beam_evidence_preserves_explicit_empty_support_constraints():
    solution = {
        "request": {
            "beam_type_label": "连续梁",
            "load_type_label": "均布荷载",
            "load_type": "uniform",
            "material_id": "q345",
            "spans": [4.0],
            "span_ids": ["L1"],
            "span_E_gpa": [206.0],
            "span_I_cm4": [85000.0],
            "E_gpa": 206.0,
            "I_cm4": 85000.0,
            "A_cm2": 120.0,
            "q_kn": 0.0,
            "total_length": 4.0,
        },
        "support_specs": [
            {"id": "A", "x": 0.0, "type": "fixed", "constraints": [], "springs": [{"dof": "rz", "stiffnessKnMPerRad": 12000.0}]},
        ],
        "support_positions": [0.0],
        "x_data": [0.0],
        "element_end_moments": [0.0],
        "element_end_shears": [0.0],
        "reactions": [],
        "max_deflection_mm": 0.0,
        "max_deflection_position_m": 0.0,
        "allowable_mm": 16.0,
    }

    boundary_text = _table_text(build_evidence_tables(solution, "beam", "测试材料")["边界条件表"])

    assert "A | x=0.0 m | 固结支座 | 无固定约束" in boundary_text
    assert "v 竖向挠度" not in boundary_text


def test_report_review_table_records_review_status_contract_source_and_diagnostics():
    table = build_report_review_table(
        {
            "resultSource": {"source": "combination", "id": "ULS1", "label": "基本组合", "description": "1.2DL + 1.4LL"},
            "benchmark": {"caseId": "BM-001", "verificationLevelLabel": "A 级验证"},
            "diagnostics": {"issues": [{"title": "支座约束不足"}, {"code": "LOAD_CASE_MISSING"}]},
        },
        "frame",
        {"reviewStatus": "ready_for_review"},
    )
    rows = table.astype(str).values.tolist()

    assert any(row[0] == "审阅状态" and row[1] == "可审阅" for row in rows)
    assert any(row[0] == "ASMS-JSON 契约版本" and row[1] == API_SCHEMA_VERSION for row in rows)
    assert any(row[0] == "结果来源" and "荷载组合: 基本组合 [ULS1]" in row[1] for row in rows)
    assert any(row[0] == "公开验证参考" and "BM-001 / A 级验证" in row[1] for row in rows)
    assert any(row[0] == "诊断警告" and "支座约束不足" in row[1] and "LOAD_CASE_MISSING" in row[1] for row in rows)


def test_report_review_table_records_result_object_revision_and_model_signatures():
    table = build_report_review_table(
        {
            "resultSource": {"source": "case", "id": "LC1", "label": "恒载", "description": "恒载工况"},
            "resultProvenance": {
                "analysisObjectId": "frame-object-1",
                "projectRevision": 12,
                "currentProjectRevision": 14,
                "modelSignature": "fnv1a64:1234567890abcdef",
                "modelHash": "backend-model-hash",
                "requestHash": "backend-request-hash",
                "solvedAt": "2026-07-19T12:00:00.000Z",
            },
        },
        "frame",
    )
    rows = table.astype(str).values.tolist()

    assert any(row[0] == "分析对象 ID" and row[1] == "frame-object-1" for row in rows)
    assert any(row[0] == "工程修订" and "计算时 12" in row[1] and "导出时 14" in row[1] for row in rows)
    assert any(row[0] == "模型签名" and "fnv1a64:1234567890abcdef" in row[1] and "backend-model-hash" in row[1] for row in rows)
    assert any(row[0] == "请求签名" and "backend-request-hash" in row[1] for row in rows)


def test_frame_export_templates_split_stability_summary_and_detail():
    payload = {
        "analysisType": "frame",
        "projectName": "稳定审查模板回归",
        "materialId": "q345",
        "structure": {
            "template": "explicit",
            "nodes": [
                {"id": "N1", "x": 0.0, "y": 0.0, "supportType": "fixed"},
                {"id": "N2", "x": 0.0, "y": 5.0, "supportType": "free"},
            ],
            "members": [
                {"id": "C1", "start": "N1", "end": "N2", "E_GPa": 210.0, "A_cm2": 220.0, "I_cm4": 1200.0, "kind": "column"},
            ],
            "loads": [
                {"type": "nodal", "node": "N2", "fxKn": 80.0, "fyKn": -300.0, "mzKnM": 0.0},
            ],
            "loadCases": [
                {
                    "id": "DL",
                    "title": "恒载",
                    "loads": [{"type": "nodal", "node": "N2", "fxKn": 0.0, "fyKn": -40.0, "mzKnM": 0.0}],
                },
                {
                    "id": "WL",
                    "title": "风载",
                    "loads": [{"type": "nodal", "node": "N2", "fxKn": 8.0, "fyKn": 0.0, "mzKnM": 0.0}],
                },
            ],
            "loadCombinations": [
                {"id": "ULS1", "title": "基本组合", "factors": {"DL": 1.2, "WL": 1.5}},
            ],
        },
        "analysisOptions": {
            "pDelta": True,
            "buckling": True,
            "pDeltaOptions": {
                "algorithm": "corotational_newton_v1",
                "initialStep": 1.0,
                "minStep": 0.2,
                "maxStep": 1.0,
                "maxIterations": 1,
                "maxCutbacks": 1,
                "includeMethodComparison": True,
                "initialImperfection": {
                    "type": "explicit",
                    "nodeOffsets": [{"nodeId": "N2", "uxMm": 6.0, "uyMm": 0.0}],
                },
            },
        },
    }

    standard_report = build_report_model(
        payload,
        analysis_type="frame",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
        report_options={"template": "standard", "figureMode": "overlay", "figureScope": "all", "reviewStatus": "draft"},
    )
    complete_report = build_report_model(
        payload,
        analysis_type="frame",
        material_name="测试材料",
        sensitivity_results=None,
        report_images=None,
        report_options={"template": "complete", "figureMode": "both", "figureScope": "all", "reviewStatus": "draft"},
    )

    standard_docx = Document(export_report(standard_report, "docx").buffer)
    standard_text = "\n".join(
        [paragraph.text for paragraph in standard_docx.paragraphs]
        + [cell.text for table in standard_docx.tables for row in table.rows for cell in row.cells]
    )
    assert "稳定审查摘要" in standard_text
    assert "P-Delta 路径控制" in standard_text
    assert "P-Delta 路径关键点" in standard_text
    assert "P-Delta 最后收敛点" in standard_text
    assert "P-Delta 失败尝试" in standard_text
    assert "初始缺陷说明" in standard_text
    assert "方法比较" in standard_text
    assert "构件 Euler K=1 初筛仅用于定位复核对象，不替代整体屈曲结论" in standard_text
    assert "corotational_newton_v1" in standard_text
    assert "initial_stress_v1" in standard_text
    assert "linear_buckling_v1" in standard_text
    assert "cutback" in standard_text
    assert "failure" in standard_text
    assert "sourceHash" in standard_text
    assert "requestHash" in standard_text
    assert "modelHash" in standard_text
    assert "residual_reduction" in standard_text
    assert "explicit" in standard_text
    assert any(
        reason in standard_text
        for reason in (
            "maximum_iterations_exhausted",
            "maximum_cutbacks_exhausted",
            "minimum_step_exhausted",
            "line_search_failed",
            "singular_tangent",
            "non_finite_increment",
        )
    )
    assert not any(reason in standard_text for reason in ("P-Delta 收敛记录", "lineSearchScale", "energyIncrementJ"))
    assert "屈曲节点模态向量" not in standard_text
    assert "屈曲构件模态形状" not in standard_text
    assert "共回转计算原理" not in standard_text

    complete_docx = Document(export_report(complete_report, "docx").buffer)
    complete_text = "\n".join(
        [paragraph.text for paragraph in complete_docx.paragraphs]
        + [cell.text for table in complete_docx.tables for row in table.rows for cell in row.cells]
    )
    assert "稳定审查摘要" in complete_text
    assert "构件 Euler K=1 初筛仅用于定位复核对象，不替代整体屈曲结论" in complete_text
    assert "稳定审查过程" in complete_text
    assert "lineSearchScale" in complete_text
    assert "energyIncrementJ" in complete_text
    assert "P-Delta 收敛记录" in complete_text
    assert "屈曲节点模态向量" in complete_text
    assert "屈曲构件模态形状" in complete_text
    assert "共回转计算原理" in complete_text
    assert "g(u, λ) = f_ext(λ) - f_int(u) = 0" in complete_text
    assert "共回转代表单元" in complete_text
    assert "当前弦长 L" in complete_text

    standard_xlsx = pd.read_excel(export_report(standard_report, "xlsx").buffer, sheet_name=None, header=None)
    standard_xlsx_text = "\n".join(frame.astype(str).to_string() for frame in standard_xlsx.values())
    assert "稳定审查摘要" in standard_xlsx_text
    assert "P-Delta 路径控制" in standard_xlsx_text
    assert "P-Delta 路径关键点" in standard_xlsx_text
    assert "P-Delta 最后收敛点" in standard_xlsx_text
    assert "P-Delta 失败尝试" in standard_xlsx_text
    assert "初始缺陷说明" in standard_xlsx_text
    assert "方法比较" in standard_xlsx_text
    assert "sourceHash" in standard_xlsx_text
    assert "requestHash" in standard_xlsx_text
    assert "modelHash" in standard_xlsx_text
    assert "P-Delta 收敛记录" not in standard_xlsx_text
    assert "屈曲节点模态向量" not in standard_xlsx_text

    complete_xlsx = pd.read_excel(export_report(complete_report, "xlsx").buffer, sheet_name=None, header=None)
    complete_xlsx_text = "\n".join(frame.astype(str).to_string() for frame in complete_xlsx.values())
    assert "稳定审查摘要" in complete_xlsx_text
    assert "lineSearchTrials" in complete_xlsx_text
    assert "energyIncrementJ" in complete_xlsx_text
    assert "P-Delta 收敛记录" in complete_xlsx_text
    assert "屈曲节点模态向量" in complete_xlsx_text
    assert "屈曲构件模态形状" in complete_xlsx_text
    assert "共回转计算原理" in complete_xlsx_text
    assert "共回转代表单元" in complete_xlsx_text


def test_standard_and_complete_exports_diverge_by_analysis_type_for_docx_and_xlsx():
    cases = [
        ("beam", beam_payload()),
        ("frame", frame_payload()),
        ("truss", truss_payload()),
    ]
    for analysis_type, payload in cases:
        standard_report = build_report_model(
            payload,
            analysis_type=analysis_type,
            material_name="测试材料",
            sensitivity_results=None,
            report_images=None,
            report_options={"template": "standard", "figureMode": "overlay", "figureScope": "all", "reviewStatus": "draft"},
        )
        complete_report = build_report_model(
            payload,
            analysis_type=analysis_type,
            material_name="测试材料",
            sensitivity_results=None,
            report_images=None,
            report_options={"template": "complete", "figureMode": "both", "figureScope": "all", "reviewStatus": "draft"},
        )

        standard_docx_text = _docx_text(export_report(standard_report, "docx"))
        complete_docx_text = _docx_text(export_report(complete_report, "docx"))
        standard_report_xlsx = build_report_model(
            payload,
            analysis_type=analysis_type,
            material_name="测试材料",
            sensitivity_results=None,
            report_images=None,
            report_options={"template": "standard", "figureMode": "overlay", "figureScope": "all", "reviewStatus": "draft"},
        )
        complete_report_xlsx = build_report_model(
            payload,
            analysis_type=analysis_type,
            material_name="测试材料",
            sensitivity_results=None,
            report_images=None,
            report_options={"template": "complete", "figureMode": "both", "figureScope": "all", "reviewStatus": "draft"},
        )
        standard_xlsx = pd.ExcelFile(io.BytesIO(export_report(standard_report_xlsx, "xlsx").buffer.getvalue()))
        complete_xlsx = pd.ExcelFile(io.BytesIO(export_report(complete_report_xlsx, "xlsx").buffer.getvalue()))

        assert "工程输入摘要" in standard_docx_text
        assert "关键点表" in standard_docx_text
        assert "CalculationTrace" not in standard_docx_text
        assert "复核点表" not in standard_docx_text
        assert "包络来源" not in standard_docx_text
        assert "计算快照" not in standard_docx_text

        assert "工程输入摘要" in complete_docx_text
        assert "关键点表" in complete_docx_text
        assert "CalculationTrace" in complete_docx_text
        assert "复核点表" in complete_docx_text
        assert "包络来源" in complete_docx_text
        assert "计算快照" in complete_docx_text

        if analysis_type == "frame":
            standard_xlsx_text = _xlsx_text(export_report(standard_report_xlsx, "xlsx"))
            complete_xlsx_text = _xlsx_text(export_report(complete_report_xlsx, "xlsx"))
            assert "P-Delta 收敛记录" not in standard_xlsx_text
            assert "P-Delta 收敛记录" in complete_xlsx_text
            assert "屈曲节点模态向量" in complete_xlsx_text
            assert "屈曲构件模态形状" in complete_xlsx_text
        else:
            assert len(complete_xlsx.sheet_names) == len(standard_xlsx.sheet_names) + 1
            assert set(complete_xlsx.sheet_names) - set(standard_xlsx.sheet_names)
