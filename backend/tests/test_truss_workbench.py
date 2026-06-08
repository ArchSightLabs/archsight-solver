import base64
import copy
import io
import os
import sys

import pytest
import pandas as pd
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

from app import app
from backend.exporters.common.report_figure_catalog import TRUSS_REPORT_OVERLAY_FIGURES, report_figures_for_scope
from backend.exporters.common.report_figures import BLUE, ChartSeries, line_chart_png
from backend.tests.docx_assertions import assert_docx_embeds_report_images
from backend.tests.benchmark_catalog import load_benchmark_catalog


BENCHMARK_CASES = [case for case in load_benchmark_catalog()["cases"] if case["category"] == "truss"]


def _report_image(seed: int = 1) -> str:
    png = line_chart_png([0, 1], [ChartSeries("占位", [0, 1], BLUE)], width=900 + seed, height=480)
    return f"data:image/png;base64,{base64.b64encode(png).decode('ascii')}"


TRUSS_DATA_CURVE_IMAGE_KEYS = ("truss.curve.ux", "truss.curve.uy", "truss.curve.axial")
TRUSS_DATA_CURVE_TITLES = ("节点 X 向位移数据曲线", "节点 Y 向位移数据曲线", "杆件轴力数据曲线")


def _truss_report_images_for_scope(*, include_all: bool, include_data_curves: bool = False) -> dict[str, str]:
    keys = [
        "truss.preview",
        *(figure.image_key for figure in report_figures_for_scope(TRUSS_REPORT_OVERLAY_FIGURES, include_all=include_all)),
        *(TRUSS_DATA_CURVE_IMAGE_KEYS if include_data_curves else ()),
    ]
    return {key: _report_image(index) for index, key in enumerate(keys, start=1)}


def _base_payload():
    return copy.deepcopy(BENCHMARK_CASES[0]["payload"])


@pytest.fixture
def client():
    app.config["TESTING"] = True
    return app.test_client()


def test_truss_benchmark_case_catalog_is_populated():
    assert len(BENCHMARK_CASES) >= 1


@pytest.mark.parametrize("case", BENCHMARK_CASES, ids=lambda case: case["id"])
def test_truss_case_regression(client, case):
    response = client.post("/api/calculate", json=case["payload"])
    assert response.status_code == 200

    data = response.get_json()
    expected = case["expected"]
    tolerances = case["tolerances"]

    assert data["analysisType"] == "truss"
    assert data["summary"]["statusCode"] == expected["statusCode"]
    assert data["summary"]["method"] == "二维平面桁架杆单元法"
    assert len(data["nodeIds"]) == expected["nodeCount"]
    assert len(data["memberIds"]) == expected["memberCount"]
    assert data["summary"]["maxDisplacementMm"] == pytest.approx(
        expected["maxDisplacementMm"], abs=tolerances["maxDisplacementMm"]
    )
    assert data["summary"]["maxAxialForceKn"] == pytest.approx(
        expected["maxAxialForceKn"], abs=tolerances["maxAxialForceKn"]
    )
    assert data["summary"]["maxDisplacementNodeId"] == expected["maxDisplacementNodeId"]
    assert data["summary"]["maxAxialForceMemberId"] == expected["maxAxialForceMemberId"]
    assert len(data["nodeResults"]) == expected["nodeCount"]
    assert len(data["memberResults"]) == expected["memberCount"]
    assert data["memberResults"][0]["maxAbsAxialForceKn"] == pytest.approx(abs(data["memberResults"][0]["axialForceKn"]))
    assert data["summary"]["peakInternalForces"]["maxAbsAxialForceKn"] == pytest.approx(data["summary"]["maxAxialForceKn"])
    assert data["summary"]["equilibriumRmsRelativeError"] < 1e-9
    assert data["diagnostics"]["equilibrium"]["rmsRelativeError"] < 1e-9


def test_truss_preview_endpoint_returns_deformed_shape(client):
    response = client.post("/api/preview", json=_base_payload())
    assert response.status_code == 200
    data = response.get_json()

    assert data["analysisType"] == "truss"
    assert data["preview"]["analysisType"] == "truss"
    assert data["preview"]["structureTypeLabel"] == "二维平面桁架"
    assert len(data["preview"]["nodes"]) == 4
    assert len(data["preview"]["members"]) == 5
    assert len(data["preview"]["loads"]) == 2
    assert len(data["preview"]["deformedNodes"]) == 4
    assert [node["supportType"] for node in data["preview"]["nodes"][:2]] == ["pinned", "roller"]
    assert data["summary"]["statusCode"] == "PASS"


def test_truss_xlsx_export_contains_core_sheets(client):
    response = client.post("/api/export", json={**_base_payload(), "format": "xlsx"})
    assert response.status_code == 200
    assert "spreadsheetml.sheet" in response.mimetype
    assert len(response.data) > 1000
    with pd.ExcelFile(io.BytesIO(response.data)) as xls:
        assert xls.sheet_names == [
            "01_复核总览",
            "02_输入模型",
            "03_单位换算",
            "04_边界条件",
            "05_校核证据",
            "06_结果明细",
            "99_原始数据",
        ]
        overview_text = pd.read_excel(xls, sheet_name="01_复核总览", header=None).to_string()
        boundary_text = pd.read_excel(xls, sheet_name="04_边界条件", header=None).to_string()
        evidence_text = pd.read_excel(xls, sheet_name="05_校核证据", header=None).to_string()
        result_text = pd.read_excel(xls, sheet_name="06_结果明细", header=None).to_string()

    assert "关键控制项" in overview_text
    assert "仅 ux/uy 平动支座约束" in boundary_text
    assert "弹性约束刚度" not in boundary_text
    assert "模型假定与适用范围" in evidence_text
    assert "施加 ux / uy 平动支座约束后求解节点位移" in evidence_text
    assert "施加支座约束与弹性约束刚度" not in evidence_text
    assert "校核证据" in evidence_text
    assert "节点结果" in result_text
    assert "杆件结果" in result_text


def test_truss_docx_export_smoke(client):
    response = client.post("/api/export", json={**_base_payload(), "format": "docx"})
    assert response.status_code == 200
    assert "wordprocessingml.document" in response.mimetype
    assert len(response.data) > 1000

    doc = Document(io.BytesIO(response.data))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = " ".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "平面桁架工程计算书" in full_text
    assert "1. 项目概况" in full_text
    assert "2. 输入参数" in full_text
    assert "2.1 受力变形图" in full_text
    assert "后端兜底示意" not in full_text
    assert "已跳过受力变形插图" in full_text
    assert "带节点编号、杆件编号、尺寸与荷载标注" in full_text
    assert "2.2 可审查计算证据链" in full_text
    assert "桁架" in full_text
    assert "仅承受轴力" in table_text
    assert "施加 ux / uy 平动支座约束后求解节点位移" in table_text
    assert "弹性约束刚度" not in table_text
    assert "施加支座约束与弹性约束刚度" not in table_text
    assert "校核证据" in full_text
    assert "3.1 节点水平位移图" not in full_text
    assert "3.2 节点竖向位移图" not in full_text
    assert "4.1 杆件轴力图" not in full_text
    assert "计算简图与结果同图显示" not in full_text
    assert "未收到前端同源工程图或数据曲线" in full_text
    assert "杆件轴力图、节点位移图" in full_text
    assert "杆件轴力曲线" not in full_text
    assert "5. 校核结论" in full_text
    assert "7. 附录数据" in full_text
    assert len(doc.inline_shapes) == 0


def test_truss_docx_export_uses_ui_overlay_figures_for_complete_scope(client):
    figures = report_figures_for_scope(TRUSS_REPORT_OVERLAY_FIGURES, include_all=True)
    report_images = _truss_report_images_for_scope(include_all=True, include_data_curves=True)

    response = client.post(
        "/api/export",
        json={
            **_base_payload(),
            "format": "docx",
            "reportOptions": {"template": "complete", "figureMode": "both", "figureScope": "all"},
            "reportImages": report_images,
        },
    )

    assert response.status_code == 200
    doc = Document(io.BytesIO(response.data))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "3.1 节点水平位移图" not in full_text
    assert "图 2-1 平面桁架受力变形示意（节点、杆件编号、尺寸与荷载标注同图显示；蓝色为放大后的变形线）" in full_text
    expected_image_keys = ["truss.preview", *(figure.image_key for figure in figures), *TRUSS_DATA_CURVE_IMAGE_KEYS]
    assert list(report_images) == expected_image_keys
    for index, figure in enumerate(figures, start=1):
        assert f"4.{index} {figure.title}" in full_text
        assert f"图 4-{index} {figure.title}（{figure.unit}，模型叠加工程图）" in full_text
    for offset, title in enumerate(TRUSS_DATA_CURVE_TITLES, start=1):
        index = len(figures) + offset
        assert f"4.{index} {title}" in full_text
        assert f"图 4-{index} {title}" in full_text
    assert "未收到前端同源受力变形图" not in full_text
    assert "未收到前端同源工程图或数据曲线" not in full_text
    assert len(doc.inline_shapes) == 1 + len(figures) + len(TRUSS_DATA_CURVE_IMAGE_KEYS)
    assert_docx_embeds_report_images(
        response.data,
        report_images,
        expected_image_keys,
    )


def test_truss_docx_export_legacy_control_scope_uses_all_core_figures(client):
    figures = report_figures_for_scope(TRUSS_REPORT_OVERLAY_FIGURES, include_all=True)
    report_images = _truss_report_images_for_scope(include_all=True)

    response = client.post(
        "/api/export",
        json={
            **_base_payload(),
            "format": "docx",
            "reportOptions": {"template": "standard", "figureMode": "overlay", "figureScope": "control"},
            "reportImages": report_images,
        },
    )

    assert response.status_code == 200
    doc = Document(io.BytesIO(response.data))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "图 2-1 平面桁架受力变形示意（节点、杆件编号、尺寸与荷载标注同图显示；蓝色为放大后的变形线）" in full_text
    for index, figure in enumerate(figures, start=1):
        assert f"4.{index} {figure.title}" in full_text
        assert f"图 4-{index} {figure.title}（{figure.unit}，模型叠加工程图）" in full_text
    assert "未收到前端同源工程图或数据曲线" not in full_text
    assert len(doc.inline_shapes) == 1 + len(figures)
    assert_docx_embeds_report_images(
        response.data,
        report_images,
        ["truss.preview", *(figure.image_key for figure in figures)],
    )


@pytest.mark.parametrize(
    ("mutate_payload", "expected_error"),
    [
        (lambda payload: payload["structure"].__setitem__("nodes", []), "桁架至少需要 2 个节点"),
        (lambda payload: payload["structure"].__setitem__("members", []), "桁架至少需要 1 个杆件"),
        (
            lambda payload: payload["structure"]["nodes"].__setitem__(1, {**payload["structure"]["nodes"][1], "id": "N1"}),
            "节点 ID 重复: N1",
        ),
        (
            lambda payload: payload["structure"]["members"].__setitem__(1, {**payload["structure"]["members"][1], "id": "M1"}),
            "构件 ID 重复: M1",
        ),
        (
            lambda payload: payload["structure"]["members"].__setitem__(0, {**payload["structure"]["members"][0], "end": "NX"}),
            "构件 M1 的起止节点无效",
        ),
        (
            lambda payload: payload["structure"]["members"].__setitem__(0, {**payload["structure"]["members"][0], "end": "N1"}),
            "构件 M1 长度必须大于 0",
        ),
        (
            lambda payload: payload["structure"]["members"].__setitem__(0, {**payload["structure"]["members"][0], "A_cm2": 0}),
            "构件 M1 的截面面积必须大于 0",
        ),
        (
            lambda payload: payload["structure"]["loads"].append({"type": "distributed", "member": "M1", "direction": "local_y", "wyKnPerM": -1.0}),
            "桁架构件荷载方向必须为 global_x 或 global_y",
        ),
        (
            lambda payload: payload["structure"]["loads"].__setitem__(0, {**payload["structure"]["loads"][0], "node": "NX"}),
            "节点荷载引用了不存在的节点",
        ),
        (
            lambda payload: payload["structure"]["members"].pop(1),
            "桁架刚度矩阵奇异，请检查支座与杆件连接",
        ),
        (
            lambda payload: payload["structure"]["nodes"].__setitem__(0, {**payload["structure"]["nodes"][0], "supportType": "free"}),
            "桁架约束条件不足，系统无稳定自由度可求解",
        ),
        (
            lambda payload: [
                payload["structure"]["nodes"].__setitem__(index, {**node, "supportType": "pinned"})
                for index, node in enumerate(payload["structure"]["nodes"])
            ],
            "桁架约束条件过多，系统无自由度可求解",
        ),
    ],
)
def test_truss_calculate_returns_stable_validation_errors(client, mutate_payload, expected_error):
    payload = _base_payload()
    mutate_payload(payload)

    response = client.post("/api/calculate", json=payload)
    assert response.status_code == 400
    data = response.get_json()
    assert data["success"] is False
    assert data["error"]["message"] == expected_error
    assert data["error"]["code"] == "TRUSS_INVALID_REQUEST"


def test_truss_member_self_weight_load_is_equivalent_nodal_load(client):
    payload = _base_payload()
    payload["structure"]["loads"] = [{"type": "distributed", "member": "M2", "selfWeightKnPerM": 2.0}]

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    equivalent_loads = data["structure"]["loads"]
    assert equivalent_loads == [
        {"type": "nodal", "node": "N3", "fxKn": 0.0, "fyKn": -2.0},
        {"type": "nodal", "node": "N4", "fxKn": 0.0, "fyKn": -2.0},
    ]
    assert data["summary"]["maxAxialForceKn"] > 0.0


def test_truss_load_cases_combinations_and_envelope(client):
    payload = _base_payload()
    payload["structure"]["loads"] = []
    payload["structure"]["loadCases"] = [
        {"id": "VL", "title": "竖向节点荷载", "loads": [{"type": "nodal", "node": "N3", "fyKn": -12.0, "fxKn": 0.0}]},
        {"id": "HL", "title": "水平节点荷载", "loads": [{"type": "nodal", "node": "N4", "fxKn": 8.0, "fyKn": 0.0}]},
    ]
    payload["structure"]["loadCombinations"] = [{"id": "COMB1", "title": "组合", "factors": {"VL": 1.2, "HL": 1.0}, "tags": ["ULS"]}]

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    assert [item["id"] for item in data["loadCaseResults"]] == ["VL", "HL"]
    assert data["loadCombinationResults"][0]["id"] == "COMB1"
    assert data["structure"]["loadCombinations"][0]["tags"] == ["ULS"]
    assert data["loadCombinationResults"][0]["tags"] == ["ULS"]
    assert data["envelope"]["maxNodeDisplacementMm"] > 0.0
    assert data["envelope"]["maxAxialForceKn"] > 0.0
    assert data["envelope"]["maxReactionKn"] > 0.0


def test_truss_exports_describe_support_nodes_and_constraints(client):
    payload = _base_payload()
    payload["structure"]["members"][0]["E_GPa"] = 30.0

    xlsx_response = client.post("/api/export", json={**payload, "format": "xlsx"})
    assert xlsx_response.status_code == 200
    with pd.ExcelFile(io.BytesIO(xlsx_response.data)) as xls:
        model_text = pd.read_excel(xls, sheet_name="02_输入模型", header=None).to_string()

    assert "支座节点" in model_text
    assert "N1：铰支座（约束 ux、uy）" in model_text
    assert "N2：滚动支座（约束 uy）" in model_text
    assert "材料适用范围" in model_text
    assert "桁架整体刚度按各杆件 E_GPa / A_cm2 输入装配" in model_text
    assert "E=30 GPa：1 个杆件" in model_text
    assert "平面桁架节点仅含 ux、uy 平动自由度" in model_text
    assert "弯矩主指标" in model_text

    docx_response = client.post("/api/export", json={**payload, "format": "docx"})
    assert docx_response.status_code == 200
    doc = Document(io.BytesIO(docx_response.data))
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "N1：铰支座（约束 ux、uy）" in table_text
    assert "N2：滚动支座（约束 uy）" in table_text
    assert "E=30 GPa：1 个杆件" in table_text
    assert "平面桁架节点仅含 ux、uy 平动自由度" in table_text
