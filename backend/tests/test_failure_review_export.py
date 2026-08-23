from __future__ import annotations

import io
import json
import os
import hashlib
import sys

import pandas as pd
import pytest
from docx import Document
from werkzeug.http import parse_options_header

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

from app import app
from backend.services.job_store import store_job


def _download_name(response) -> str:
    _, options = parse_options_header(response.headers.get("Content-Disposition", ""))
    return options.get("filename", "")


def _failure_review_payload(format_type: str = "docx") -> dict:
    return {
        "format": format_type,
        "inputId": "case-20260823-001",
        "completedStages": ["输入校验", "结构装配", "稳定分析"],
        "stableErrorCode": "FRAME_PDELTA_NOT_CONVERGED",
        "objectRefs": [
            {"kind": "member", "id": "B1"},
            {"kind": "node", "id": "N4"},
        ],
        "diagnostics": [
            {
                "code": "FRAME_PDELTA_NOT_CONVERGED",
                "title": "P-Delta 未收敛",
                "detail": "二阶迭代在给定荷载步内未达到收敛容差。",
                "severity": "error",
            },
            {
                "code": "STRUCTURE_UNSTABLE_CONSTRAINTS",
                "title": "结构约束不足",
                "detail": "侧向约束不足，体系无法形成稳定二阶路径。",
                "severity": "warning",
            },
        ],
        "hashes": {
            "requestHash": "req-8b3b1b3f",
            "modelHash": "mdl-72f44e0d",
            "snapshotHash": "snap-4f5cc69e",
        },
        "suggestedActions": [
            "增加 loadSteps 并提高 maxIterations。",
            "补充侧向约束或检查支座释放。",
        ],
    }


def _failed_job_payload() -> dict:
    return {
        "analysisType": "frame",
        "projectName": "failed-job-frame",
        "materialId": "Q345B",
    }


def _stable_hash(value) -> str:
    serialized = json.dumps(value, ensure_ascii=False, separators=(",", ":"), sort_keys=True)
    return hashlib.sha256(serialized.encode("utf-8")).hexdigest()


def _docx_text(response) -> str:
    doc = Document(io.BytesIO(response.data))
    return "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )


@pytest.fixture
def client():
    app.config["TESTING"] = True
    return app.test_client()


def test_failure_review_export_route_returns_docx_and_omits_result_fields(client):
    response = client.post("/api/export/failure", json=_failure_review_payload("docx"))

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert _download_name(response) == "case_20260823_001_失败审查_审查材料.docx"

    text = _docx_text(response)
    assert "失败审查材料" in text
    assert "case-20260823-001" in text
    assert "FRAME_PDELTA_NOT_CONVERGED" in text
    assert "输入校验" in text
    assert "结构装配" in text
    assert "稳定分析" in text
    assert "req-8b3b1b3f" in text
    assert "mdl-72f44e0d" in text
    assert "snap-4f5cc69e" in text
    assert "增加 loadSteps 并提高 maxIterations。" in text
    assert "补充侧向约束或检查支座释放。" in text
    for forbidden in ("位移", "反力", "内力", "关键点", "包络", "计算完成"):
        assert forbidden not in text


def test_failure_review_export_material_type_route_returns_xlsx(client):
    payload = {**_failure_review_payload("xlsx"), "materialType": "failure-review"}
    response = client.post("/api/export", json=payload)

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    with pd.ExcelFile(io.BytesIO(response.data)) as xls:
        assert xls.sheet_names == ["01_失败审查总览", "02_对象定位", "03_诊断记录", "04_建议动作"]
        overview = pd.read_excel(xls, sheet_name="01_失败审查总览", header=None).to_string()
        objects = pd.read_excel(xls, sheet_name="02_对象定位", header=None).to_string()
        diagnostics = pd.read_excel(xls, sheet_name="03_诊断记录", header=None).to_string()
        actions = pd.read_excel(xls, sheet_name="04_建议动作", header=None).to_string()

    assert "失败审查总览" in overview
    assert "case-20260823-001" in overview
    assert "FRAME_PDELTA_NOT_CONVERGED" in overview
    assert "B1" in objects
    assert "N4" in objects
    assert "P-Delta 未收敛" in diagnostics
    assert "结构约束不足" in diagnostics
    assert "增加 loadSteps 并提高 maxIterations。" in actions
    assert "补充侧向约束或检查支座释放。" in actions
    for forbidden in ("位移", "反力", "内力", "关键点", "包络", "计算完成"):
        assert forbidden not in (overview + objects + diagnostics + actions)


def test_failure_review_export_reads_failed_job_from_job_store(client, monkeypatch, tmp_path):
    monkeypatch.setenv("ARCHSIGHT_SOLVER_JOB_DB_PATH", str(tmp_path / "solver-jobs.sqlite3"))
    now = "2026-08-23T08:00:00+00:00"
    payload = _failed_job_payload()
    store_job(
        {
            "jobId": "failed-job-001",
            "clientJobId": "failed-client-001",
            "operation": "calculate",
            "payload": payload,
            "status": "failed",
            "error": {
                "code": "STRUCTURE_UNSTABLE_CONSTRAINTS",
                "message": "约束条件不足，系统无稳定自由度可求解。",
            },
            "warnings": ["约束条件不足，系统无稳定自由度可求解。"],
            "infos": ["输入已归一化", "求解器装配完成"],
            "createdAt": now,
            "updatedAt": now,
            "startedAt": now,
            "completedAt": now,
        }
    )

    response = client.post(
        "/api/export/failure",
        json={
            "jobId": "failed-job-001",
            "format": "docx",
        },
    )

    assert response.status_code == 200
    text = _docx_text(response)
    assert "failed-client-001" in text
    assert "STRUCTURE_UNSTABLE_CONSTRAINTS" in text
    assert "结构约束不足" in text
    assert _stable_hash(payload) in text
    assert "求解器装配完成" not in text
    for forbidden in ("位移", "反力", "内力", "关键点", "包络", "计算完成"):
        assert forbidden not in text

    compat_response = client.post(
        "/api/export",
        json={
            "materialType": "failure-review",
            "jobId": "failed-job-001",
            "format": "docx",
        },
    )
    assert compat_response.status_code == 200


def test_failure_review_export_job_mode_rejects_manual_fabrication(client, monkeypatch, tmp_path):
    monkeypatch.setenv("ARCHSIGHT_SOLVER_JOB_DB_PATH", str(tmp_path / "solver-jobs.sqlite3"))
    now = "2026-08-23T08:00:00+00:00"
    store_job(
        {
            "jobId": "failed-job-002",
            "operation": "calculate",
            "payload": _failed_job_payload(),
            "status": "failed",
            "error": {
                "code": "COMMON_ASYNC_JOB_FAILED",
                "message": "异步求解作业失败: 人工伪造数据。",
            },
            "warnings": [],
            "infos": [],
            "createdAt": now,
            "updatedAt": now,
            "startedAt": now,
            "completedAt": now,
        }
    )

    response = client.post(
        "/api/export/failure",
        json={
            "jobId": "failed-job-002",
            "format": "docx",
            "inputId": "fake-input",
            "completedStages": ["fake"],
            "stableErrorCode": "FAKE_CODE",
            "diagnostics": [],
            "hashes": {},
            "suggestedActions": [],
        },
    )

    assert response.status_code == 400
    data = response.get_json()
    assert data["error"]["code"] == "FAILURE_REVIEW_INVALID_REQUEST"
    assert "jobId 模式不接受" in data["error"]["message"]


def test_failure_review_export_rejects_result_like_fields(client):
    response = client.post(
        "/api/export/failure",
        json={
            **_failure_review_payload("docx"),
            "criticalPoints": [{"name": "控制点", "value": 1.0}],
            "governingEnvelope": {"maxDisplacementMm": 12.3},
            "calculationTrace": {"steps": []},
        },
    )

    assert response.status_code == 400
    data = response.get_json()
    assert data["success"] is False
    assert data["error"]["code"] == "FAILURE_REVIEW_INVALID_REQUEST"
    assert "不允许" in data["error"]["message"]
