import copy
import io
import os
import sys

import pandas as pd
import pytest
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

from app import app


@pytest.fixture
def client():
    app.config["TESTING"] = True
    return app.test_client()


def beam_payload():
    return {
        "analysisType": "beam",
        "projectName": "Span Property Regression",
        "materialId": "q345",
        "beamType": "continuous",
        "loadType": "uniform",
        "spans": [5.0, 5.0],
        "spanProperties": [
            {"E": 210.0, "I": 4500.0},
            {"E": 210.0, "I": 4500.0},
        ],
        "q": 10.0,
        "freq": 2.0,
        "duration": 5.0,
    }


def test_beam_calculate_response_preserves_span_properties(client):
    response = client.post("/api/calculate", json=beam_payload())
    assert response.status_code == 200

    data = response.get_json()

    assert data["analysisType"] == "beam"
    assert data["payload"]["spanProperties"] == beam_payload()["spanProperties"]
    assert len(data["payload"]["spanProperties"]) == len(data["payload"]["spans"])


def test_beam_second_span_property_changes_result(client):
    baseline_response = client.post("/api/calculate", json=beam_payload())
    assert baseline_response.status_code == 200
    baseline = baseline_response.get_json()

    variant = copy.deepcopy(beam_payload())
    variant["spanProperties"][1]["I"] = 450.0

    variant_response = client.post("/api/calculate", json=variant)
    assert variant_response.status_code == 200
    changed = variant_response.get_json()

    assert changed["payload"]["spanProperties"][1]["I"] == 450.0
    assert changed["summary"]["maxDeflectionMm"] > baseline["summary"]["maxDeflectionMm"] * 1.05


def test_beam_supports_query_points_and_timoshenko_option(client):
    payload = beam_payload()
    payload.update(
        {
            "spans": [4.0],
            "beamType": "continuous",
            "loadType": "uniform",
            "q": 12.0,
            "E": 30.0,
            "G": 12.0,
            "A_cm2": 300.0,
            "I": 1200.0,
            "beamTheory": "timoshenko",
            "supports": [
                {"id": "A", "x": 0.0, "type": "pinned"},
                {"id": "B", "x": 4.0, "type": "roller"},
            ],
            "queryPointsM": [0.0, 2.0, 4.0],
        }
    )

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    assert data["payload"]["beamTheory"] == "timoshenko"
    assert data["beam"]["beamTheoryLabel"] == "Timoshenko 梁理论"
    assert [item["label"] for item in data["beam"]["supports"]] == ["A", "B"]
    assert [item["supportId"] for item in data["beam"]["reactions"]] == ["A", "B"]
    assert [item["xM"] for item in data["queryResults"]] == [0.0, 2.0, 4.0]
    assert abs(data["queryResults"][1]["momentKnM"]) > 0.0


def test_beam_query_point_ratios_and_symbolic_check(client):
    payload = beam_payload()
    payload.update(
        {
            "spans": [6.0],
            "beamType": "simply_supported",
            "loadType": "uniform",
            "q": 10.0,
            "queryPointRatios": [0.0, 0.5, 1.0],
        }
    )

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    assert [item["xM"] for item in data["queryResults"]] == [0.0, 3.0, 6.0]
    assert data["symbolicCheck"]["available"] is True
    assert data["symbolicCheck"]["equations"] == ["R_A = R_B = qL / 2", "M_max = qL^2 / 8", "v_max = 5qL^4 / (384EI)"]
    assert data["summary"]["maxPositiveMomentKnM"] >= data["summary"]["maxNegativeMomentKnM"]


def test_uniform_load_moment_diagram_uses_smooth_samples(client):
    payload = beam_payload()
    payload.update(
        {
            "spans": [10.0],
            "beamType": "simply_supported",
            "loadType": "uniform",
            "q": 10.0,
        }
    )

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    assert data["x_data"][1] > 0.0
    assert data["moment_data"][1] > 0.0
    assert data["summary"]["maxPositiveMomentKnM"] == pytest.approx(125.0, rel=1e-3)


def test_partial_uniform_load_preview_respects_range(client):
    payload = beam_payload()
    payload.update(
        {
            "spans": [8.0],
            "beamType": "simply_supported",
            "loadType": "uniform",
            "loads": [{"type": "uniform", "qKnPerM": 5.0, "start": 2.0, "end": 6.0}],
        }
    )

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    assert data["normalizedRequest"]["loads"][0]["uniform_start"] == 2.0
    assert data["normalizedRequest"]["loads"][0]["uniform_end"] == 6.0
    assert data["diagram"]["loadItems"][0]["start"] == 2.0
    assert data["diagram"]["loadItems"][0]["end"] == 6.0
    assert data["beam"]["loads"] == [
        {
            "type": "uniform",
            "x": 4.0,
            "intensityKnPerM": 5.0,
            "startX": 2.0,
            "endX": 6.0,
            "length": 4.0,
        }
    ]


def test_linear_load_preview_respects_range_ratios(client):
    payload = beam_payload()
    payload.update(
        {
            "spans": [4.0],
            "beamType": "continuous",
            "loadType": "linear",
            "distributedLoadStart": 10.0,
            "distributedLoadEnd": 10.7,
            "distributedLoadStartRatio": 0.0,
            "distributedLoadEndRatio": 0.05,
        }
    )

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    loads = data["beam"]["loads"]
    assert [round(item["x"], 6) for item in loads] == [0.0, 0.2]
    assert loads[0]["startX"] == 0.0
    assert loads[0]["endX"] == 0.2
    assert data["normalizedRequest"]["distributed_start_ratio"] == 0.0
    assert data["normalizedRequest"]["distributed_end_ratio"] == 0.05


def test_beam_combined_loads_support_multiple_point_loads(client):
    payload = beam_payload()
    payload.update(
        {
            "spans": [4.0, 4.0],
            "beamType": "continuous",
            "loads": [
                {"type": "uniform", "qKnPerM": 4.0},
                {"type": "linear", "qStartKnPerM": 2.0, "qEndKnPerM": 6.0, "start": 1.0, "end": 7.0},
                {"type": "point", "pointLoadKn": 8.0, "x": 2.0},
                {"type": "point", "pointLoadKn": 5.0, "x": 6.0},
            ],
        }
    )

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    assert data["payload"]["loadType"] == "combined"
    assert data["beam"]["loadTypeLabel"] == "组合荷载"
    assert [load["type"] for load in data["normalizedRequest"]["loads"]] == ["uniform", "linear", "point", "point"]
    assert [load["type"] for load in data["diagram"]["loadItems"]] == ["uniform", "linear", "point", "point"]
    assert len([load for load in data["beam"]["loads"] if load["type"] == "point"]) == 2
    assert data["summary"]["maxDeflectionMm"] > 0.0


def test_beam_elastic_support_reduces_deflection(client):
    soft = beam_payload()
    soft.update(
        {
            "spans": [6.0],
            "beamType": "continuous",
            "supports": [
                {"id": "A", "x": 0.0, "type": "pinned"},
                {"id": "M", "x": 3.0, "type": "free", "springs": [{"dof": "v", "stiffnessKnPerM": 500.0}]},
                {"id": "B", "x": 6.0, "type": "roller"},
            ],
        }
    )
    stiff = copy.deepcopy(soft)
    stiff["supports"][1]["springs"][0]["stiffnessKnPerM"] = 50000.0

    soft_data = client.post("/api/calculate", json=soft).get_json()
    stiff_data = client.post("/api/calculate", json=stiff).get_json()

    assert stiff_data["summary"]["maxDeflectionMm"] < soft_data["summary"]["maxDeflectionMm"]
    assert any(abs(item.get("springVertical", 0.0)) > 0 for item in stiff_data["solution"]["reactions"])


def test_beam_support_constraints_override_nominal_type(client):
    pinned = beam_payload()
    pinned.update(
        {
            "spans": [4.0],
            "beamType": "continuous",
            "supports": [
                {"id": "A", "x": 0.0, "type": "pinned"},
                {"id": "B", "x": 4.0, "type": "roller"},
            ],
        }
    )
    fixed = copy.deepcopy(pinned)
    fixed["supports"][0]["constraints"] = ["v", "rz"]

    pinned_data = client.post("/api/calculate", json=pinned).get_json()
    fixed_data = client.post("/api/calculate", json=fixed).get_json()

    assert fixed_data["payload"]["supports"][0]["constraints"] == ["v", "rz"]
    assert fixed_data["beam"]["supports"][0]["constraints"] == ["v", "rz"]
    assert fixed_data["summary"]["maxDeflectionMm"] < pinned_data["summary"]["maxDeflectionMm"]


def test_beam_exports_describe_material_scope_span_stiffness_and_supports(client):
    payload = beam_payload()
    payload.update(
        {
            "projectName": "Beam Semantic Export",
            "spans": [3.0, 5.0],
            "spanProperties": [
                {"id": "L1", "E": 210.0, "I": 4500.0},
                {"id": "L2", "E": 30.0, "I": 900.0},
            ],
            "supports": [
                {"id": "A", "x": 0.0, "type": "fixed"},
                {"id": "M", "x": 3.0, "type": "free", "springs": [{"dof": "v", "stiffnessKnPerM": 800.0}]},
                {"id": "B", "x": 8.0, "type": "roller"},
            ],
        }
    )

    xlsx_response = client.post("/api/export", json={**payload, "format": "xlsx"})
    assert xlsx_response.status_code == 200
    with pd.ExcelFile(io.BytesIO(xlsx_response.data)) as xls:
        model_text = pd.read_excel(xls, sheet_name="02_输入模型", header=None).to_string()
        boundary_text = pd.read_excel(xls, sheet_name="04_边界条件", header=None).to_string()

    assert "材料适用范围" in model_text
    assert "梁系整体刚度按各跨段 E_GPa / I_cm4 输入装配" in model_text
    assert "跨段刚度输入" in model_text
    assert "L2" in model_text
    assert "30.0 GPa" in model_text
    assert "900.0 cm^4" in model_text
    assert "支座体系说明" in model_text
    assert "A" in boundary_text
    assert "固结支座" in boundary_text
    assert "M" in boundary_text
    assert "v 竖向挠度: 800.0 kN/m" in boundary_text
    assert "左支座" not in model_text
    assert "右支座" not in model_text

    docx_response = client.post("/api/export", json={**payload, "format": "docx"})
    assert docx_response.status_code == 200
    doc = Document(io.BytesIO(docx_response.data))
    paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "图 2-1 梁系受力变形示意（支座、荷载、跨段编号与放大后的挠度形态同图显示）" in paragraph_text
    assert "梁体受力变形" not in paragraph_text
    assert "材料适用范围" in table_text
    assert "梁系整体刚度按各跨段 E_GPa / I_cm4 输入装配" in table_text
    assert "跨段刚度输入" in paragraph_text
    assert "L2" in table_text
    assert "30.0 GPa" in table_text
    assert "900.0 cm^4" in table_text
    assert "A" in table_text
    assert "固结支座" in table_text
    assert "v 竖向挠度: 800.0 kN/m" in table_text
    assert "左支座" not in table_text
    assert "右支座" not in table_text


def test_beam_load_cases_combinations_and_envelope(client):
    payload = beam_payload()
    payload["loadCases"] = [
        {"id": "DL", "title": "恒载", "loads": [{"type": "uniform", "qKnPerM": 8.0}]},
        {"id": "LL", "title": "活载", "loads": [{"type": "point", "magnitudeKn": 20.0, "x": 2.5}]},
    ]
    payload["loadCombinations"] = [{"id": "ULS1", "title": "基本组合", "factors": {"DL": 1.2, "LL": 1.4}, "tags": ["ULS", "包络"]}]

    response = client.post("/api/calculate", json=payload)

    assert response.status_code == 200
    data = response.get_json()
    assert [item["id"] for item in data["loadCaseResults"]] == ["DL", "LL"]
    assert data["loadCombinationResults"][0]["id"] == "ULS1"
    assert data["normalizedRequest"]["loadCombinations"][0]["tags"] == ["ULS", "包络"]
    assert data["loadCombinationResults"][0]["tags"] == ["ULS", "包络"]
    assert data["envelope"]["maxDeflectionMm"] >= data["loadCaseResults"][0]["summary"]["maxDeflectionMm"]
    assert "maxPositiveMomentKnM" in data["envelope"]
    assert "maxNegativeMomentKnM" in data["envelope"]
