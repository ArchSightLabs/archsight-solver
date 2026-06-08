import base64
import io
import os
import sys

import pandas as pd
import pytest
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

from app import app
from backend.exporters.common.report_figure_catalog import FRAME_REPORT_MEMBER_FIGURES, report_figures_for_scope
from backend.exporters.common.report_figures import BLUE, ChartSeries, line_chart_png
from backend.tests.docx_assertions import assert_docx_embeds_report_images


def _report_image(seed: int = 1) -> str:
    png = line_chart_png([0, 1], [ChartSeries("占位", [0, 1], BLUE)], width=900 + seed, height=480)
    return f"data:image/png;base64,{base64.b64encode(png).decode('ascii')}"


FRAME_DATA_CURVE_IMAGE_KEYS = ("frame.curve.ux", "frame.curve.uy")
FRAME_DATA_CURVE_TITLES = ("节点 X 向位移数据曲线", "节点 Y 向位移数据曲线")


def _frame_report_images_for_scope(*, include_all: bool, include_data_curves: bool = False) -> dict[str, str]:
    keys = [
        "frame.preview",
        *(figure.overlay_image_key for figure in report_figures_for_scope(FRAME_REPORT_MEMBER_FIGURES, include_all=include_all)),
        *(FRAME_DATA_CURVE_IMAGE_KEYS if include_data_curves else ()),
    ]
    return {key: _report_image(index) for index, key in enumerate(keys, start=1)}


def frame_payload():
    return {
        "analysisType": "frame",
        "projectName": "Portal Frame Test",
        "materialId": "q345",
        "structure": {
            "template": "portal_frame",
            "span": 6.0,
            "height": 4.0,
            "left_support": "fixed",
            "right_support": "fixed",
            "beam_load_kn_per_m": 18.0,
            "lateral_load_kn": 24.0,
            "top_vertical_load_kn": 0.0,
            "nodes": [
                {"id": "N1", "x": 0.0, "y": 0.0, "supportType": "fixed"},
                {"id": "N2", "x": 6.0, "y": 0.0, "supportType": "fixed"},
                {"id": "N3", "x": 0.0, "y": 4.0, "supportType": "free"},
                {"id": "N4", "x": 6.0, "y": 4.0, "supportType": "free"},
            ],
            "members": [
                {"id": "C1", "start": "N1", "end": "N3", "E_GPa": 210, "A_cm2": 240, "I_cm4": 12000, "kind": "column"},
                {"id": "B1", "start": "N3", "end": "N4", "E_GPa": 210, "A_cm2": 220, "I_cm4": 15000, "kind": "beam"},
                {"id": "C2", "start": "N2", "end": "N4", "E_GPa": 210, "A_cm2": 240, "I_cm4": 12000, "kind": "column"},
            ],
            "loads": [
                {"type": "distributed", "member": "B1", "wyKnPerM": -18.0},
                {"type": "nodal", "node": "N4", "fxKn": 24.0, "fyKn": 0.0, "mzKnM": 0.0},
            ],
        },
    }


def explicit_frame_payload():
    return {
        "analysisType": "frame",
        "projectName": "Explicit Frame Test",
        "materialId": "q345",
        "structure": {
            "template": "explicit",
            "nodes": [
                {"id": "N1", "x": 0.0, "y": 0.0, "supportType": "fixed"},
                {"id": "N2", "x": 4.0, "y": 0.0, "supportType": "pinned"},
                {"id": "N3", "x": 8.0, "y": 0.0, "supportType": "roller"},
                {"id": "N4", "x": 0.0, "y": 4.0, "supportType": "free"},
                {"id": "N5", "x": 4.0, "y": 4.0, "supportType": "free"},
                {"id": "N6", "x": 8.0, "y": 4.0, "supportType": "free"},
            ],
            "members": [
                {"id": "C1", "start": "N1", "end": "N4", "E_GPa": 210, "A_cm2": 260, "I_cm4": 14000, "kind": "column"},
                {"id": "C2", "start": "N2", "end": "N5", "E_GPa": 210, "A_cm2": 260, "I_cm4": 14000, "kind": "column"},
                {"id": "C3", "start": "N3", "end": "N6", "E_GPa": 210, "A_cm2": 260, "I_cm4": 14000, "kind": "column"},
                {"id": "B1", "start": "N4", "end": "N5", "E_GPa": 210, "A_cm2": 220, "I_cm4": 15000, "kind": "beam"},
                {"id": "B2", "start": "N5", "end": "N6", "E_GPa": 210, "A_cm2": 220, "I_cm4": 15000, "kind": "beam"},
            ],
            "loads": [
                {"type": "distributed", "member": "B1", "wyKnPerM": -14.0},
                {"type": "distributed", "member": "B2", "wyKnPerM": -10.0},
                {"type": "nodal", "node": "N6", "fxKn": 16.0, "fyKn": 0.0, "mzKnM": 0.0},
            ],
        },
    }


def frame_payload_with_combination_tags():
    payload = frame_payload()
    payload["projectName"] = "Portal Frame Load Combination Tags"
    payload["structure"]["loads"] = []
    payload["structure"]["loadCases"] = [
        {
            "id": "DL",
            "title": "恒载",
            "loads": [{"type": "distributed", "member": "B1", "wyKnPerM": -18.0}],
        },
        {
            "id": "WL",
            "title": "风载",
            "loads": [{"type": "nodal", "node": "N4", "fxKn": 24.0, "fyKn": 0.0, "mzKnM": 0.0}],
        },
    ]
    payload["structure"]["loadCombinations"] = [
        {"id": "ULS1", "title": "承载能力基本组合", "factors": {"DL": 1.2, "WL": 1.5}, "tags": ["ULS", "包络"]},
    ]
    return payload


EXPLICIT_TWO_BAY_EXPECTED = {
    "node_count": 6,
    "member_count": 5,
    "method": "二维平面框架刚度法 + 平面梁柱单元",
    "expected_upgrade_condition": "取得手算、教材或可信工具复核值后升级为数值基准",
}


@pytest.fixture
def client():
    app.config["TESTING"] = True
    return app.test_client()


@pytest.fixture
def explicit_two_bay_frame_payload():
    return explicit_frame_payload()


def test_frame_calculate_returns_frame_results(client):
    response = client.post("/api/calculate", json=frame_payload())
    assert response.status_code == 200
    data = response.get_json()

    assert data["analysisType"] == "frame"
    assert "frame" in data
    assert "summary" in data
    assert len(data["nodeIds"]) == 4
    assert len(data["memberIds"]) == 3
    assert len(data["ux_data"]) == 4
    assert len(data["uy_data"]) == 4
    assert len(data["member_moment_data"]) == 3
    assert data["summary"]["maxDisplacementMm"] > 0
    assert data["summary"]["status"] in ("合格", "需校核")
    assert data["summary"]["method"] == "二维平面框架刚度法 + 平面梁柱单元"
    assert data["preview"]["structureTypeLabel"] == "二维平面框架"
    assert data["diagram"]["structureTypeLabel"] == "二维平面框架"


def test_frame_distributed_load_negative_wy_deflects_downward(client):
    payload = frame_payload()
    payload["structure"]["loads"][1]["fxKn"] = 0.0

    response = client.post("/api/calculate", json=payload)
    assert response.status_code == 200
    data = response.get_json()
    node_results = {item["nodeId"]: item for item in data["nodeResults"]}

    assert node_results["N3"]["uyMm"] < 0
    assert node_results["N4"]["uyMm"] < 0


def test_frame_template_request_respects_snake_case_contract_fields(client):
    response = client.post(
        "/api/calculate",
        json={
            "analysisType": "frame",
            "projectName": "Template Contract Test",
            "materialId": "q345",
            "structure": {
                "template": "portal_frame",
                "span": 10.0,
                "height": 5.0,
                "left_support": "pinned",
                "right_support": "fixed",
                "beam_load_kn_per_m": 5.0,
                "lateral_load_kn": 7.0,
                "top_vertical_load_kn": 3.0,
                "columnE_GPa": 200,
                "beamE_GPa": 190,
                "columnA_cm2": 300,
                "beamA_cm2": 250,
                "columnI_cm4": 20000,
                "beamI_cm4": 25000,
            },
        },
    )
    assert response.status_code == 200
    structure = response.get_json()["structure"]

    assert structure["span"] == 10.0
    assert structure["height"] == 5.0
    assert structure["left_support"] == "pinned"
    assert structure["right_support"] == "fixed"
    assert structure["beam_load_kn_per_m"] == 5.0
    assert structure["lateral_load_kn"] == 7.0
    assert structure["top_vertical_load_kn"] == 3.0
    assert structure["loads"] == [
        {"type": "distributed", "member": "B1", "wyKnPerM": -5.0},
        {"type": "nodal", "node": "N4", "fxKn": 7.0, "fyKn": -3.0, "mzKnM": 0.0},
    ]


def test_frame_preview_endpoint_returns_deformed_shape(client):
    response = client.post("/api/preview", json=frame_payload())
    assert response.status_code == 200
    data = response.get_json()

    assert data["analysisType"] == "frame"
    assert "frame" in data
    assert len(data["frame"]["nodes"]) == 4
    assert len(data["frame"]["deformedNodes"]) == 4
    assert len(data["frame"]["loads"]) == 2


def test_frame_xlsx_export_contains_core_sheets(client):
    response = client.post("/api/export", json={**frame_payload(), "format": "xlsx"})
    assert response.status_code == 200
    assert "spreadsheetml.sheet" in response.mimetype

    excel_data = io.BytesIO(response.data)
    with pd.ExcelFile(excel_data) as xls:
        assert xls.sheet_names == [
            "01_复核总览",
            "02_输入模型",
            "03_单位换算",
            "04_边界条件",
            "05_校核证据",
            "06_结果明细",
            "99_原始数据",
        ]
        overview_text = pd.read_excel(xls, sheet_name="01_复核总览", header=None).to_string()
        model_text = pd.read_excel(xls, sheet_name="02_输入模型", header=None).to_string()
        evidence_text = pd.read_excel(xls, sheet_name="05_校核证据", header=None).to_string()
        result_text = pd.read_excel(xls, sheet_name="06_结果明细", header=None).to_string()
        raw_text = pd.read_excel(xls, sheet_name="99_原始数据", header=None).to_string()

    assert "关键控制项" in overview_text
    assert "二维平面框架" in overview_text
    assert "二维平面框架" in model_text
    assert "节点模型" in model_text
    assert "构件模型" in model_text
    assert "模型假定与适用范围" in evidence_text
    assert "校核证据" in evidence_text
    assert "节点结果" in result_text
    assert "构件结果" in result_text
    assert "构件内力曲线" in raw_text


def test_frame_docx_export_smoke(client):
    response = client.post("/api/export", json={**frame_payload(), "format": "docx"})

    assert response.status_code == 200
    assert "wordprocessingml.document" in response.mimetype
    assert len(response.data) > 1000

    doc = Document(io.BytesIO(response.data))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "平面框架工程计算书" in full_text
    assert "1. 项目概况" in full_text
    assert "2. 输入参数" in full_text
    assert "2.1 受力变形图" in full_text
    assert "后端兜底示意" not in full_text
    assert "已跳过受力变形插图" in full_text
    assert "带节点编号、构件编号、尺寸与荷载标注" in full_text
    assert "2.2 可审查计算证据链" in full_text
    assert "模型假定与适用范围" in full_text
    assert "边界条件表" in full_text
    assert "校核证据" in full_text
    assert "3.1 节点水平位移图" not in full_text
    assert "3.2 节点竖向位移图" not in full_text
    assert "4.1 构件弯矩图（模型叠加）" not in full_text
    assert "计算简图与结果同图显示" not in full_text
    assert "未收到前端同源工程图或数据曲线" in full_text
    assert "构件弯矩图、构件剪力图、构件局部 y 向挠度图、构件轴力图" in full_text
    assert "4.2 构件剪力" not in full_text
    assert "构件弯矩曲线" not in full_text
    assert "5. 校核结论" in full_text
    assert "7. 附录数据" in full_text
    assert len(doc.inline_shapes) == 0


def test_frame_docx_export_uses_ui_overlay_figures_for_complete_scope(client):
    figures = report_figures_for_scope(FRAME_REPORT_MEMBER_FIGURES, include_all=True)
    report_images = _frame_report_images_for_scope(include_all=True, include_data_curves=True)

    response = client.post(
        "/api/export",
        json={
            **frame_payload(),
            "format": "docx",
            "reportOptions": {"template": "complete", "figureMode": "both", "figureScope": "all"},
            "reportImages": report_images,
        },
    )

    assert response.status_code == 200
    doc = Document(io.BytesIO(response.data))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "3.1 节点水平位移图" not in full_text
    assert "构件弯矩曲线" not in full_text
    assert "构件剪力曲线" not in full_text
    assert "图 2-1 平面框架受力变形示意（节点、构件编号、尺寸与荷载标注同图显示；蓝色为放大后的变形线）" in full_text
    expected_image_keys = ["frame.preview", *(figure.overlay_image_key for figure in figures), *FRAME_DATA_CURVE_IMAGE_KEYS]
    assert list(report_images) == expected_image_keys
    for index, figure in enumerate(figures, start=1):
        assert f"4.{index} 构件{figure.title}（模型叠加）" in full_text
        assert f"图 4-{index} 构件{figure.title}（{figure.unit}，模型叠加工程图）" in full_text
    for offset, title in enumerate(FRAME_DATA_CURVE_TITLES, start=1):
        index = len(figures) + offset
        assert f"4.{index} {title}" in full_text
        assert f"图 4-{index} {title}" in full_text
    assert "未收到前端同源受力变形图" not in full_text
    assert "未收到前端同源工程图或数据曲线" not in full_text
    assert len(doc.inline_shapes) == 1 + len(figures) + len(FRAME_DATA_CURVE_IMAGE_KEYS)
    assert_docx_embeds_report_images(
        response.data,
        report_images,
        expected_image_keys,
    )


def test_frame_docx_export_legacy_control_scope_uses_all_core_figures(client):
    figures = report_figures_for_scope(FRAME_REPORT_MEMBER_FIGURES, include_all=True)
    report_images = _frame_report_images_for_scope(include_all=True)

    response = client.post(
        "/api/export",
        json={
            **frame_payload(),
            "format": "docx",
            "reportOptions": {"template": "standard", "figureMode": "overlay", "figureScope": "control"},
            "reportImages": report_images,
        },
    )

    assert response.status_code == 200
    doc = Document(io.BytesIO(response.data))
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "图 2-1 平面框架受力变形示意（节点、构件编号、尺寸与荷载标注同图显示；蓝色为放大后的变形线）" in full_text
    for index, figure in enumerate(figures, start=1):
        assert f"4.{index} 构件{figure.title}（模型叠加）" in full_text
        assert f"图 4-{index} 构件{figure.title}（{figure.unit}，模型叠加工程图）" in full_text
    assert "未收到前端同源工程图或数据曲线" not in full_text
    assert len(doc.inline_shapes) == 1 + len(figures)
    assert_docx_embeds_report_images(
        response.data,
        report_images,
        ["frame.preview", *(figure.overlay_image_key for figure in figures)],
    )


def test_frame_exports_include_load_combination_tags(client):
    xlsx_response = client.post("/api/export", json={**frame_payload_with_combination_tags(), "format": "xlsx"})
    assert xlsx_response.status_code == 200
    with pd.ExcelFile(io.BytesIO(xlsx_response.data)) as xls:
        raw_text = pd.read_excel(xls, sheet_name="99_原始数据", header=None).to_string()

    assert "ULS" in raw_text
    assert "包络" in raw_text

    docx_response = client.post("/api/export", json={**frame_payload_with_combination_tags(), "format": "docx"})
    assert docx_response.status_code == 200
    doc = Document(io.BytesIO(docx_response.data))
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "ULS" in table_text
    assert "包络" in table_text


def test_frame_exports_describe_support_nodes_instead_of_left_right(client):
    payload = explicit_frame_payload()
    payload["structure"]["nodes"][2]["supportAngleDeg"] = 45.0
    payload["structure"]["members"][1]["E_GPa"] = 30.0

    xlsx_response = client.post("/api/export", json={**payload, "format": "xlsx"})
    assert xlsx_response.status_code == 200
    with pd.ExcelFile(io.BytesIO(xlsx_response.data)) as xls:
        model_text = pd.read_excel(xls, sheet_name="02_输入模型", header=None).to_string()
        boundary_text = pd.read_excel(xls, sheet_name="04_边界条件", header=None).to_string()

    assert "支座节点" in model_text
    assert "N1：固结支座" in model_text
    assert "N2：铰支座" in model_text
    assert "N3：滚动支座，法向角 45°" in model_text
    assert "法向位移 n（45°）" in boundary_text
    assert "释放切向位移 t" in boundary_text
    assert "释放 rz 平面转角" in boundary_text
    assert "材料适用范围" in model_text
    assert "框架整体刚度按各构件 E_GPa / A_cm2 / I_cm4 输入装配" in model_text
    assert "E=30 GPa：1 个构件" in model_text
    assert "框架支座按节点 ux / uy / rz 自由度参与整体刚度矩阵" in model_text
    assert "左支座" not in model_text
    assert "右支座" not in model_text

    docx_response = client.post("/api/export", json={**payload, "format": "docx"})
    assert docx_response.status_code == 200
    doc = Document(io.BytesIO(docx_response.data))
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "N3：滚动支座，法向角 45°" in table_text
    assert "法向位移 n（45°）" in table_text
    assert "释放切向位移 t" in table_text
    assert "释放 rz 平面转角" in table_text
    assert "E=30 GPa：1 个构件" in table_text
    assert "左支座" not in table_text
    assert "右支座" not in table_text


def test_frame_calculate_supports_explicit_two_bay_fixture(client, explicit_two_bay_frame_payload):
    response = client.post("/api/calculate", json=explicit_two_bay_frame_payload)
    assert response.status_code == 200
    data = response.get_json()

    assert data["analysisType"] == "frame"
    assert data["summary"]["method"] == EXPLICIT_TWO_BAY_EXPECTED["method"]
    assert len(data["nodeIds"]) == EXPLICIT_TWO_BAY_EXPECTED["node_count"]
    assert len(data["memberIds"]) == EXPLICIT_TWO_BAY_EXPECTED["member_count"]
    assert data["summary"]["maxDisplacementMm"] > 0
    assert data["summary"]["status"] in ("合格", "需校核")


@pytest.mark.parametrize(
    ("mutate_payload", "expected_error"),
    [
        (
            lambda payload: payload["structure"]["nodes"].__setitem__(1, {**payload["structure"]["nodes"][1], "id": "N1"}),
            "节点 ID 重复: N1",
        ),
        (
            lambda payload: payload["structure"]["members"].__setitem__(1, {**payload["structure"]["members"][1], "id": "C1"}),
            "构件 ID 重复: C1",
        ),
        (
            lambda payload: payload["structure"]["members"].__setitem__(0, {**payload["structure"]["members"][0], "end": "NX"}),
            "构件 C1 的起止节点无效",
        ),
        (
            lambda payload: payload["structure"]["loads"].append({"type": "nodal", "node": "NX", "fxKn": 1.0}),
            "节点荷载引用了不存在的节点",
        ),
        (
            lambda payload: payload["structure"]["loads"].append({"type": "distributed", "member": "MX", "wyKnPerM": -1.0}),
            "构件荷载引用了不存在的构件",
        ),
        (
            lambda payload: payload["structure"]["members"].__setitem__(0, {**payload["structure"]["members"][0], "end": "N1"}),
            "构件 C1 长度必须大于 0",
        ),
    ],
)
def test_frame_calculate_returns_stable_validation_errors(client, explicit_two_bay_frame_payload, mutate_payload, expected_error):
    mutate_payload(explicit_two_bay_frame_payload)

    response = client.post("/api/calculate", json=explicit_two_bay_frame_payload)

    assert response.status_code == 400
    data = response.get_json()
    assert data["success"] is False
    assert data["error"]["message"] == expected_error
    assert data["error"]["code"] == "FRAME_INVALID_REQUEST"
